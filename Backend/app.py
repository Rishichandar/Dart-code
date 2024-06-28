



# from flask import Flask, request, jsonify, render_template, redirect, url_for
# from flask_sqlalchemy import SQLAlchemy
# from flask_cors import CORS
# from werkzeug.utils import secure_filename
# import pandas as pd
# import os
# import json
# import bcrypt
# import io
# import docx
# import numpy as np
# from sklearn.impute import SimpleImputer
# from sklearn.linear_model import LinearRegression
# from sklearn.ensemble import RandomForestRegressor
# from sklearn.preprocessing import PowerTransformer, StandardScaler, MinMaxScaler
# from scipy.stats import boxcox
# from io import StringIO
# from sklearn.model_selection import train_test_split
# from sklearn.decomposition import PCA
# from sklearn.preprocessing import MinMaxScaler, PolynomialFeatures
# from sklearn.feature_selection import SelectKBest, f_classif
# from sklearn.decomposition import PCA
# from sklearn.impute import SimpleImputer
# import base64
# from sklearn.ensemble import RandomForestClassifier
# import xgboost as xgb
# import matplotlib.pyplot as plt

# app = Flask(__name__)
# CORS(app)

# # Database configuration
# app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+mysqlconnector://root:rishi123@localhost:3306/dart'
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# db = SQLAlchemy(app)

# # Define the User model
# class User(db.Model):
#     __tablename__ = 'dart_users'
#     userid = db.Column(db.Integer, primary_key=True)
#     email = db.Column(db.String(120), unique=True, nullable=False)
#     password = db.Column(db.String(120), nullable=False)
#     username = db.Column(db.String(120), nullable=False)
#     role = db.Column(db.String(120), nullable=False)

#     def __init__(self, email, password, username, role):
#         self.email = email
#         self.password = password
#         self.username = username
#         self.role = role

# # Create the database tables
# with app.app_context():
#     db.create_all()

# # User signup endpoint
# @app.route('/user/signup', methods=['POST'])
# def signup():
#     data = request.get_json()
#     email = data.get('email')
#     password = data.get('password')
#     username = data.get('username')
#     role = data.get('role')  # Get role from request

#     if not email or not password or not username or not role:
#         return jsonify({'error': 'Email, password, username, and role are required'}), 400

#     # Check if the user already exists
#     existing_user = User.query.filter_by(email=email).first()
#     if existing_user:
#         return jsonify({'error': 'User already exists'}), 400

#     # Hash the password
#     hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

#     new_user = User(email=email, password=hashed_password, username=username, role=role)
#     db.session.add(new_user)
#     db.session.commit()

#     # Generate a token (you can replace this with actual JWT generation)
#     token = 'some_generated_token'  # Replace with real token generation logic

#     return jsonify({'token': token}), 200

# # User login endpoint
# @app.route('/user/login', methods=['POST'])
# def login():
#     data = request.get_json()
#     email = data.get('email')
#     password = data.get('password')

#     if not email or not password:
#         return jsonify({'error': 'Email and password are required'}), 400

#     # Check if the user exists
#     user = User.query.filter_by(email=email).first()
#     if not user:
#         return jsonify({'error': 'Invalid email or password'}), 401

#     # Verify the password
#     if not bcrypt.checkpw(password.encode('utf-8'), user.password.encode('utf-8')):
#         return jsonify({'error': 'Invalid email or password'}), 401

#     return jsonify({'message': 'success'}), 200

# # User password reset endpoint
# @app.route('/user/reset-password', methods=['POST'])
# def reset_password():
#     data = request.get_json()
#     email = data.get('email')
#     new_password = data.get('newPassword')

#     if not email or not new_password:
#         return jsonify({'error': 'Email and new password are required'}), 400

#     # Find the user by email
#     user = User.query.filter_by(email=email).first()
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Hash the new password
#     hashed_password = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
#     # Update the user's password
#     user.password = hashed_password
#     db.session.commit()

#     return jsonify({'message': 'Password reset successful'}), 200

# # Get username endpoint
# @app.route('/user/username', methods=['POST'])
# def get_username():
#     data = request.get_json()
#     email = data.get('email')

#     if not email:
#         return jsonify({'error': 'Email is required'}), 400

#     # Find the user by email
#     user = User.query.filter_by(email=email).first()
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     return jsonify({'username': user.username, 'role': user.role}), 200

# # Data preprocessing class
# class preprocess:
#     @staticmethod
#     def poz_read_file(file_path_or_uploaded_file):
#         try:
#             # Check if the input is a file path or an uploaded file object
#             if isinstance(file_path_or_uploaded_file, str):
#                 # It's a file path
#                 file_extension = os.path.splitext(file_path_or_uploaded_file)[-1].lower()
#                 with open(file_path_or_uploaded_file, 'rb') as f:
#                     file_content = f.read()
#             else:
#                 # It's an uploaded file object
#                 file_extension = file_path_or_uploaded_file.filename.split('.')[-1].lower()
#                 file_content = file_path_or_uploaded_file.read()

#             # Debug: Print file extension and size
#             print(f"File extension: {file_extension}")
#             print(f"File content size: {len(file_content)} bytes")

#             # Read the file content based on its extension
#             if file_extension == 'csv':
#                 return pd.read_csv(io.BytesIO(file_content))
#             elif file_extension in ['xls', 'xlsx']:
#                 return pd.read_excel(io.BytesIO(file_content))
#             elif file_extension == 'txt':
#                 data = file_content.decode("utf-8")
#                 return pd.DataFrame({'text_data': [data]})
#             elif file_extension in ['doc', 'docx']:
#                 doc = docx.Document(io.BytesIO(file_content))
#                 text = ' '.join([para.text for para in doc.paragraphs])
#                 return pd.DataFrame({'text_data': [text]})
#             elif file_extension == 'json':
#                 json_data = json.loads(file_content)
#                 text = ' '.join(str(value) for value in json_data.values())
#                 return pd.DataFrame({'text_data': [text]})
#             else:
#                 print("Unsupported file type.")
#                 return None
#         except Exception as e:
#             # Debug: Print exception message
#             print(f"Exception in poz_read_file: {e}")
#             return None

#     @staticmethod
#     def poz_handle_missing_values(df, fill_method='mean', constant=0):
#         try:
#             if fill_method == 'mean':
#                 fill_value = df.mean()
#                 df_filled = df.fillna(fill_value)
#             elif fill_method == 'median':
#                 fill_value = df.median()    
#                 df_filled = df.fillna(fill_value)
#             elif fill_method == 'mode':
#                 fill_value = df.mode().iloc[0]
#                 df_filled = df.fillna(fill_value)
#             elif fill_method == 'constant':
#                 df_filled = df.fillna(constant)
#             elif fill_method == 'linear':
#                 df_filled = preprocess.fill_missing_values_linear(df)
#             elif fill_method == 'random_forest':
#                 df_filled = preprocess.fill_missing_values_rf(df)
#             else:
#                 raise ValueError("Invalid fill_method. Options: 'mean', 'median', 'mode', 'constant', 'linear', 'random_forest'.")
#         except Exception as e:
#             print(f"An error occurred: {e}")
#             df_filled = pd.DataFrame()  # Return an empty DataFrame in case of error

#         return df_filled  

#     @staticmethod
#     def fill_missing_values_linear(df):
#         df_filled = df.copy()
#         initial_imputer = SimpleImputer(strategy='mean')
#         df_initial_imputed = pd.DataFrame(initial_imputer.fit_transform(df_filled), columns=df.columns)
#         for column in df.columns:
#             missing_values_mask = df[column].isnull()
#             if not missing_values_mask.any():
#                 continue
#             X_train = df_initial_imputed.loc[~missing_values_mask].drop(columns=[column])
#             y_train = df_initial_imputed.loc[~missing_values_mask, column]
#             X_test = df_initial_imputed.loc[missing_values_mask].drop(columns=[column])
#             model = LinearRegression()
#             model.fit(X_train, y_train)
#             predicted_values = model.predict(X_test)
#             df_filled.loc[missing_values_mask, column] = predicted_values
#         return df_filled

#     @staticmethod
#     def fill_missing_values_rf(df):
#         df_filled = df.copy()
#         initial_imputer = SimpleImputer(strategy='mean')
#         df_initial_imputed = pd.DataFrame(initial_imputer.fit_transform(df_filled), columns=df.columns)
#         for column in df.columns:
#             missing_values_mask = df[column].isnull()
#             if not missing_values_mask.any():
#                 continue
#             X_train = df_initial_imputed.loc[~missing_values_mask].drop(columns=[column])
#             y_train = df_initial_imputed.loc[~missing_values_mask, column]
#             X_test = df_initial_imputed.loc[missing_values_mask].drop(columns=[column])
#             model = RandomForestRegressor()
#             model.fit(X_train, y_train)
#             predicted_values = model.predict(X_test)
#             df_filled.loc[missing_values_mask, column] = predicted_values
#         return df_filled

#     @staticmethod
#     def poz_transform(df, method='standardize'):
#         try:
#             numerical_columns = df.select_dtypes(include='number')
#             if method == 'standardize':
#                 transformed_df = (numerical_columns - numerical_columns.mean()) / numerical_columns.std()
#             elif method == 'boxcox':
#                 transformed_df = numerical_columns.apply(lambda x: pd.Series(boxcox(x + 1)[0], index=x.index))
#             elif method == 'yeojohnson':
#                 pt = PowerTransformer(method='yeo-johnson')
#                 transformed_data = pt.fit_transform(numerical_columns)
#                 transformed_df = pd.DataFrame(transformed_data, columns=numerical_columns.columns, index=numerical_columns.index)
#             elif method == 'normalize':
#                 transformed_df = numerical_columns.apply(lambda x: (x - x.min()) / (x.max() - x.min()))
#             elif method == 'scaler':
#                 scaler = StandardScaler()
#                 transformed_data = scaler.fit_transform(numerical_columns)
#                 transformed_df = pd.DataFrame(transformed_data, columns=numerical_columns.columns, index=numerical_columns.index)
#             elif method == 'minmax':
#                 scaler = MinMaxScaler()
#                 transformed_data = scaler.fit_transform(numerical_columns)
#                 transformed_df = pd.DataFrame(transformed_data, columns=numerical_columns.columns, index=numerical_columns.index)
#             elif method == 'log2':
#                 transformed_df = np.log2(numerical_columns)
#             else:
#                 raise ValueError(
#                     "Invalid method. Options: 'standardize', 'boxcox', 'yeojohnson', 'normalize', 'scaler', 'minmax', 'log2'.")

#             # Combine transformed numerical columns with non-numeric columns
#             for column in df.columns:
#                 if column in numerical_columns.columns:
#                     df[column] = transformed_df[column]
#         except Exception as e:
#             print(f"An error occurred: {e}")
#             df = pd.DataFrame()  # Return an empty DataFrame in case of error
#         return df
    
#     def poz_outliers(data, method='z-score', threshold=3, action='remove'):
#         cleaned_data = data.copy()
#         if method == 'z-score':
#             # Detect outliers using z-score
#             z_scores = (cleaned_data - cleaned_data.mean()) / cleaned_data.std()
#             outliers = abs(z_scores) > threshold

#         elif method == 'iqr':
#             # Detect outliers using interquartile range (IQR)
#             q1 = cleaned_data.quantile(0.25)
#             q3 = cleaned_data.quantile(0.75)
#             iqr = q3 - q1
#             lower_bound = q1 - threshold * iqr
#             upper_bound = q3 + threshold * iqr
#             outliers = (cleaned_data < lower_bound) | (cleaned_data > upper_bound)

#         else:
#             raise ValueError("Invalid method. Options: 'z-score', 'iqr'")

#         if action == 'remove':
#             # Remove outliers
#             cleaned_data = cleaned_data[~outliers.any(axis=1)]

#         elif action == 'cap':
#             # Cap outliers to a specified range
#             cleaned_data[outliers] = cleaned_data.clip(lower=lower_bound, upper=upper_bound, axis=1)

#         else:
#             raise ValueError("Invalid action. Options: 'remove', 'transform', 'cap'")

#         return cleaned_data
    
#     @staticmethod
#     def poz_data_aggregation(data, target_column):
#         """
#         Fill missing values in the target column with the mean of that column.

#         Parameters:
#         - data: DataFrame containing the dataset.
#         - target_column: The column to fill missing values with its mean.

#         Returns:
#         - DataFrame with filled missing values in the target column.
#         """
#         try:
#             # Fill missing values in the target column with the mean of that column
#             if target_column in data.columns and pd.api.types.is_numeric_dtype(data[target_column]):
#                 mean_value = data[target_column].mean()
#                 data[target_column].fillna(mean_value, inplace=True)
            
#             return data
#         except Exception as e:
#             print(f"An error occurred during data aggregation: {e}")
#             return None
    
#     @staticmethod
#     def poz_split_data(data, target_column, test_size, random_state, axis):
#             """
#             Split data into train and test sets.
#             Parameters:
#             - data: DataFrame containing the dataset.
#             - target_column: Name of the target column.
#             - test_size: Size of the test set.
#             - random_state: Random state for reproducibility.
#             - axis: Axis along which to drop the target column (0 for rows, 1 for columns).
#             Returns:
#             - X_train, X_test, y_train, y_test: Train and test sets for features and target.
#             """
#             if axis not in [0, 1]:
#                 raise ValueError("Axis must be 0 or 1.")

#             X = data.drop(target_column, axis=axis)
#             y = data[target_column]
#             X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=random_state)
#             return X_train, X_test, y_train, y_test
   
    
#     @staticmethod
#     def poz_feature_scaling(df, scaling_method='normalization'):
#      try:
#         df_scaled = df.copy()
        
#         if scaling_method == 'normalization':
#             df_scaled = df_scaled.apply(lambda x: (x - x.min()) / (x.max() - x.min()), axis=0)
#         elif scaling_method == 'standard':
#             scaler = StandardScaler()
#             df_scaled = pd.DataFrame(scaler.fit_transform(df_scaled), columns=df_scaled.columns)
#         elif scaling_method == 'cliplog':
#             df_scaled = df_scaled.apply(lambda x: np.log1p(np.clip(x, a_min=1e-9, a_max=None)), axis=0)
#         elif scaling_method == 'z-score':
#             df_scaled = (df_scaled - df_scaled.mean()) / df_scaled.std()
#         else:
#             raise ValueError("Invalid scaling_method. Options: 'normalization', 'standard', 'cliplog', 'z-score'.")
        
#      except Exception as e:
#         print(f"An error occurred during feature scaling: {e}")
#         df_scaled = pd.DataFrame()  # Return an empty DataFrame in case of error
    
#      return df_scaled
     
#     @staticmethod
#     def feature_extraction(X, feature_names, y=None, method='pCA', n_components=2, k=5):
#         # Impute missing values
#         imputer = SimpleImputer(strategy='mean')
#         X = imputer.fit_transform(X)
        
#         if method == 'pCA':
#             pca = PCA(n_components=n_components)
#             X_transformed = pca.fit_transform(X)
#             columns = [f"PC{i + 1}" for i in range(n_components)]
#         elif method == 'scaling':
#             scaler = MinMaxScaler()
#             X_transformed = scaler.fit_transform(X)
#             columns = feature_names
#         elif method == 'polynomial':
#             poly = PolynomialFeatures(degree=2, include_bias=False)
#             X_transformed = poly.fit_transform(X)
#             poly_feature_names = poly.get_feature_names_out(input_features=feature_names)
#             columns = [name.replace(" ", "*") for name in poly_feature_names]
#         elif method == 'select k-best':
#             selector = SelectKBest(score_func=f_classif, k=k)
#             X_transformed = selector.fit_transform(X, y)
#             selected_indices = selector.get_support(indices=True)
#             columns = [feature_names[i] for i in selected_indices]
#         elif method == 'expansion':
#             poly = PolynomialFeatures(degree=3, include_bias=False)
#             X_transformed = poly.fit_transform(X)
#             poly_feature_names = poly.get_feature_names_out(input_features=feature_names)
#             columns = [name.replace(" ", "*") for name in poly_feature_names]
#         else:
#             raise ValueError(
#                 "Invalid method. Choose one of: 'pCA', 'scaling', 'polynomial', 'select k-best', 'expansion'")
#         return pd.DataFrame(X_transformed, columns=columns)
    
    
#     @staticmethod
#     def plot_feature_importances(X, y, feature_names, algorithm='Random Forest', n_estimators_rf=100, n_estimators_xgb=100, random_state=42):
#         if algorithm == 'Random Forest':
#             clf = RandomForestClassifier(n_estimators=n_estimators_rf, random_state=random_state)
#         elif algorithm == 'XGBoost':
#             clf = xgb.XGBClassifier(n_estimators=n_estimators_xgb, random_state=random_state)
#         else:
#             return "Invalid algorithm selected.", None

#         clf.fit(X, y)
#         importances = clf.feature_importances_
#         indices = np.argsort(importances)[::-1]

#         feature_ranking = []
#         for f in range(X.shape[1]):
#             feature_ranking.append((f + 1, feature_names[indices[f]], importances[indices[f]]))

#         plt.figure(figsize=(10, 6))
#         plt.title("Feature importances")
#         plt.bar(range(X.shape[1]), importances[indices], align="center")
#         plt.xticks(range(X.shape[1]), [feature_names[i] for i in indices], rotation=90)
#         plt.xlabel("Feature")
#         plt.ylabel("Feature importance")

#         img = io.BytesIO()
#         plt.savefig(img, format='png')
#         img.seek(0)
#         plot_url = base64.b64encode(img.getvalue()).decode('utf8')
#         plt.close()

#         return feature_ranking, plot_url
    
# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/upload', methods=['POST'])
# def upload():
#     if 'file' not in request.files:
#         return 'No file part'
#     file = request.files['file']
#     if file.filename == '':
#         return 'No selected file'
#     df = pd.read_csv(file)
#     feature_names = df.columns[:-1]
#     target_column = df.columns[-1]

#     X = df[feature_names].values
#     y = df[target_column].values

#     algorithm = request.form.get('algorithm', 'Random Forest')
#     n_estimators_rf = int(request.form.get('n_estimators_rf', 100))
#     n_estimators_xgb = int(request.form.get('n_estimators_xgb', 100))

#     feature_ranking, plot_url = preprocess.plot_feature_importances(X, y, feature_names, algorithm, n_estimators_rf, n_estimators_xgb)

#     return render_template('result.html', feature_ranking=feature_ranking, plot_url=plot_url)

# @app.route('/feature-importance', methods=['POST'])
# def feature_importance():
#     if 'file' not in request.files:
#         return jsonify({"error": "No file part"})
#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({"error": "No selected file"})

#     df = pd.read_csv(file)
#     target_column = request.form.get('targetColumn')
#     algorithm = request.form.get('algorithm', 'Random Forest')
#     n_estimators_rf = int(request.form.get('n_estimators_rf', 100))
#     n_estimators_xgb = int(request.form.get('n_estimators_xgb', 100))

#     # Ensure there are no NaN values in the target column
#     df = df.dropna(subset=[target_column])

#     # Convert the target column to integers if necessary
#     df[target_column] = df[target_column].astype(int)

#     feature_names = df.columns[df.columns != target_column]
#     X = df[feature_names].values
#     y = df[target_column].values

#     feature_ranking, plot_url = preprocess.plot_feature_importances(X, y, feature_names, algorithm, n_estimators_rf, n_estimators_xgb)

#     return jsonify({"feature_ranking": feature_ranking, "plot_url": plot_url})

# @app.route('/feature-extraction', methods=['POST'])
# def extract_features():
#     if 'file' not in request.files:
#         return jsonify({'error': 'No file part in the request'}), 400
#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({'error': 'No selected file'}), 400

#     method = request.form.get('method', 'pCA')
#     n_components = int(request.form.get('n_components', 2))
#     k = int(request.form.get('k', 5))

#     df = pd.read_csv(file)
#     X = df.iloc[:, :-1].values  # Assuming the last column is the target
#     feature_names = df.columns[:-1].tolist()
#     y = df.iloc[:, -1].values if method == 'select k-best' else None

#     try:
#         result = preprocess.feature_extraction(X, feature_names, y, method, n_components, k)
#         return result.to_json(orient='split')
#     except ValueError as e:
#         return jsonify({'error': str(e)}), 400


 
# # Endpoint to handle missing values
# @app.route('/handle_missing_values', methods=['POST'])
# def handle_missing_values():
#     try:
#         file = request.files['file']
#         fill_method = request.form.get('method', 'mean')
#         constant = request.form.get('constant', 0)
#         df = preprocess.poz_read_file(file)

#         if df is None:
#             raise ValueError("Error reading the file. Unsupported file type or corrupt file.")

#         df_filled = preprocess.poz_handle_missing_values(df, fill_method, constant)
#         result = df_filled.to_json(orient="split")
#         print(json.loads(result))
#         return jsonify({'processed_data': json.loads(result)}), 200
        
#     except Exception as e:
#         print(f"Error processing the DataFrame: {e}")
#         return jsonify({'error': f"Error processing the DataFrame: {e}"}), 500

# # Endpoint to handle data transformation
# @app.route('/transform_data', methods=['POST'])
# def transform_data():
#     try:
#         file = request.files['file']
#         method = request.form.get('method', 'standardize')
#         df = preprocess.poz_read_file(file)

#         if df is None:
#             raise ValueError("Error reading the file. Unsupported file type or corrupt file.")

#         df_transformed = preprocess.poz_transform(df, method)
#         result = df_transformed.to_json(orient="split")
#         print(json.loads(result))
#         return jsonify({'transformed_data': json.loads(result)}), 200
        
#     except Exception as e:
#         print(f"Error transforming the DataFrame: {e}")
#         return jsonify({'error': f"Error transforming the DataFrame: {e}"}), 500

# # Endpoint to upload CSV file
# @app.route('/upload-csv', methods=['POST'])
# def upload_csv():
#     if 'file' not in request.files:
#         return jsonify({'error': 'No file selected'}), 400
#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({'error': 'No selected file'}), 400
#     if file:
#         try:
#             # Read CSV file using pandas
#             df = pd.read_csv(file)

#             # Convert DataFrame to JSON
#             json_data = df.to_json(orient='split')

#             # Return JSON response
#             return jsonify({'data': json_data})
#         except Exception as e:
#             return jsonify({'error': str(e)}), 500
        
# @app.route('/aggregate_data', methods=['POST'])
# def handle_aggregate_data():
#     try:
#         # Check if the POST request has the file part
#         if 'file' not in request.files:
#             return jsonify({'error': 'No file part'}), 400

#         file = request.files['file']
#         target_column = request.form.get('targetColumn', None)

#         # Load CSV data from file
#         if file.filename == '':
#             return jsonify({'error': 'No selected file'}), 400

#         if file and target_column:
#             # Read CSV data
#             csv_string = file.stream.read().decode('utf-8')
#             csv_data = StringIO(csv_string)

#             # Load CSV into DataFrame
#             df = pd.read_csv(csv_data)

#             # Perform data aggregation using the predefined function
#             aggregated_data = preprocess.poz_data_aggregation(df, target_column)

#             if aggregated_data is not None:
#                 # Convert aggregated data to JSON format
#                 # aggregated_data_json = aggregated_data.to_json(orient='records')
#                 # return jsonify({'aggregated_data': aggregated_data_json}), 200
#                 aggregated_data_json = aggregated_data.to_json(orient='split')
#                 return jsonify({'aggregated_data':json.loads(aggregated_data_json)}), 200
#         #      result = df_scaled.to_json(orient="split")
#         # return jsonify({'scaled_data': json.loads(result)}), 200
#             else:
#                 return jsonify({'error': 'Data aggregation failed'}), 500

#         return jsonify({'error': 'Missing target column'}), 400

#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

        
#         # Endpoint to handle outliers
# @app.route('/handle_outliers', methods=['POST'])
# def handle_outliers():
#     try:
#         file = request.files['file']
#         method = request.form.get('method', 'z-score')
#         threshold = float(request.form.get('threshold', 3))
#         action = request.form.get('action', 'remove')
#         df = preprocess.poz_read_file(file)

#         if df is None:
#             raise ValueError("Error reading the file. Unsupported file type or corrupt file.")

#         df_outliers_handled = preprocess.poz_outliers(df, method, threshold, action)
#         result = df_outliers_handled.to_json(orient="split")
#         print(json.loads(result))
#         return jsonify({'processed_data': json.loads(result)}), 200
        
#     except Exception as e:
#         print(f"Error processing outliers: {e}")
#         return jsonify({'error': f"Error processing outliers: {e}"}), 500
    
#         # Endpoint to handle data segmentation
# @app.route('/segment_data', methods=['POST'])
# def segment_data():
#     try:
#         # Check if file is present in the request
#         if 'file' not in request.files:
#             return jsonify({'error': 'No file uploaded'}), 400

#         file = request.files['file']

#         # Read the CSV file
#         segmented_data = preprocess.poz_read_file(file)

#         # Get the last column name as the target column
#         target_column = segmented_data.columns[-1]

#         # Split the data
#         X_train, X_test, y_train, y_test = preprocess.poz_split_data(segmented_data, target_column=target_column, test_size=0.2, random_state=42, axis=1)

#         # Convert split data to JSON
#         split_data_json = {
#             'X_train': X_train.to_json(orient='split'),
#             'X_test': X_test.to_json(orient='split'),
#             'y_train': y_train.to_frame(name=target_column).to_json(orient='split'),
#             'y_test': y_test.to_frame(name=target_column).to_json(orient='split'),
#             'X_train_count':len(X_train),
#             'X_test_count':len(X_test),
#             'y_train_count': len(y_train),
#             'y_test_count': len(y_test)
#         }

#         return jsonify({'split_data': split_data_json}), 200
#     except Exception as e:
#         print(f"Error segmenting the data: {e}")
#         return jsonify({'error': f"Error segmenting the data: {e}"}), 500
    
#     # Scale features in the dataset

# # @app.route('/scale', methods=['POST'])
# # def scale():
# #     try:
# #         data = request.get_json()
# #         df = pd.DataFrame(data['data'])
# #         scaling_method = data.get('scaling_method', 'normalization')
# #         df_scaled = preprocess.poz_feature_scaling(df, scaling_method)
# #         return jsonify(df_scaled.to_dict(orient='records'))
# #     except Exception as e:
# #         return jsonify({'error': str(e)}), 400
# @app.route('/scale', methods=['POST'])
# def scale():
#     try:
#         # Retrieve CSV file from FormData
#         file = request.files['file']
        
#         # Read CSV file into pandas DataFrame
#         df = pd.read_csv(file)
        
#         # Retrieve scaling method from FormData
#         scaling_method = request.form.get('scaling_method', 'normalization')
        
#         # Perform feature scaling based on the chosen method
#         df_scaled = preprocess.poz_feature_scaling(df, scaling_method)
        
#         # Convert scaled DataFrame to JSON and return
#         # return jsonify(df_scaled.to_json(orient='records'))
#         result = df_scaled.to_json(orient="split")
#         return jsonify({'scaled_data': json.loads(result)}), 200
    
#     except Exception as e:
#         return jsonify({'error': str(e)}), 400

# if __name__ == '__main__':
#     app.run(debug=True)







from flask import Flask, request, jsonify, render_template, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from werkzeug.utils import secure_filename
import pandas as pd
import os
import json
import bcrypt
import io
import docx
import numpy as np
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.preprocessing import PowerTransformer, StandardScaler, MinMaxScaler
from scipy.stats import boxcox
from io import StringIO
from sklearn.model_selection import train_test_split
from sklearn.decomposition import PCA
from sklearn.preprocessing import MinMaxScaler, PolynomialFeatures
from sklearn.feature_selection import SelectKBest, f_classif
from sklearn.decomposition import PCA
from sklearn.impute import SimpleImputer
import base64
from sklearn.ensemble import RandomForestClassifier
import xgboost as xgb
import matplotlib.pyplot as plt

app = Flask(__name__)
CORS(app)
cors = CORS(app, resources={r"/feature-importance": {"origins": "http://localhost:3000"}})

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+mysqlconnector://root:rishi123@localhost:3306/dart'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Define the User model
class User(db.Model):
    __tablename__ = 'dart_users'
    userid = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    username = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(120), nullable=False)

    def __init__(self, email, password, username, role):
        self.email = email
        self.password = password
        self.username = username
        self.role = role

# Create the database tables
with app.app_context():
    db.create_all()

# User signup endpoint
@app.route('/user/signup', methods=['POST'])
def signup():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    username = data.get('username')
    role = data.get('role')  # Get role from request

    if not email or not password or not username or not role:
        return jsonify({'error': 'Email, password, username, and role are required'}), 400

    # Check if the user already exists
    existing_user = User.query.filter_by(email=email).first()
    if existing_user:
        return jsonify({'error': 'User already exists'}), 400

    # Hash the password
    hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

    new_user = User(email=email, password=hashed_password, username=username, role=role)
    db.session.add(new_user)
    db.session.commit()

    # Generate a token (you can replace this with actual JWT generation)
    token = 'some_generated_token'  # Replace with real token generation logic

    return jsonify({'token': token}), 200

# User login endpoint
@app.route('/user/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')

    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400

    # Check if the user exists
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': 'Invalid email or password'}), 401

    # Verify the password
    if not bcrypt.checkpw(password.encode('utf-8'), user.password.encode('utf-8')):
        return jsonify({'error': 'Invalid email or password'}), 401

    return jsonify({'message': 'success'}), 200

# User password reset endpoint
@app.route('/user/reset-password', methods=['POST'])
def reset_password():
    data = request.get_json()
    email = data.get('email')
    new_password = data.get('newPassword')

    if not email or not new_password:
        return jsonify({'error': 'Email and new password are required'}), 400

    # Find the user by email
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    # Hash the new password
    hashed_password = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    # Update the user's password
    user.password = hashed_password
    db.session.commit()

    return jsonify({'message': 'Password reset successful'}), 200

# Get username endpoint
@app.route('/user/username', methods=['POST'])
def get_username():
    data = request.get_json()
    email = data.get('email')

    if not email:
        return jsonify({'error': 'Email is required'}), 400

    # Find the user by email
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    return jsonify({'username': user.username, 'role': user.role}), 200

# Data preprocessing class
class preprocess:
    @staticmethod
    def poz_read_file(file_path_or_uploaded_file):
        try:
            # Check if the input is a file path or an uploaded file object
            if isinstance(file_path_or_uploaded_file, str):
                # It's a file path
                file_extension = os.path.splitext(file_path_or_uploaded_file)[-1].lower()
                with open(file_path_or_uploaded_file, 'rb') as f:
                    file_content = f.read()
            else:
                # It's an uploaded file object
                file_extension = file_path_or_uploaded_file.filename.split('.')[-1].lower()
                file_content = file_path_or_uploaded_file.read()

            # Debug: Print file extension and size
            print(f"File extension: {file_extension}")
            print(f"File content size: {len(file_content)} bytes")

            # Read the file content based on its extension
            if file_extension == 'csv':
                return pd.read_csv(io.BytesIO(file_content))
            elif file_extension in ['xls', 'xlsx']:
                return pd.read_excel(io.BytesIO(file_content))
            elif file_extension == 'txt':
                data = file_content.decode("utf-8")
                return pd.DataFrame({'text_data': [data]})
            elif file_extension in ['doc', 'docx']:
                doc = docx.Document(io.BytesIO(file_content))
                text = ' '.join([para.text for para in doc.paragraphs])
                return pd.DataFrame({'text_data': [text]})
            elif file_extension == 'json':
                json_data = json.loads(file_content)
                text = ' '.join(str(value) for value in json_data.values())
                return pd.DataFrame({'text_data': [text]})
            else:
                print("Unsupported file type.")
                return None
        except Exception as e:
            # Debug: Print exception message
            print(f"Exception in poz_read_file: {e}")
            return None

    @staticmethod
    def poz_handle_missing_values(df, fill_method='mean', constant=0):
        try:
            if fill_method == 'mean':
                fill_value = df.mean()
                df_filled = df.fillna(fill_value)
            elif fill_method == 'median':
                fill_value = df.median()    
                df_filled = df.fillna(fill_value)
            elif fill_method == 'mode':
                fill_value = df.mode().iloc[0]
                df_filled = df.fillna(fill_value)
            elif fill_method == 'constant':
                df_filled = df.fillna(constant)
            elif fill_method == 'linear':
                df_filled = preprocess.fill_missing_values_linear(df)
            elif fill_method == 'random_forest':
                df_filled = preprocess.fill_missing_values_rf(df)
            else:
                raise ValueError("Invalid fill_method. Options: 'mean', 'median', 'mode', 'constant', 'linear', 'random_forest'.")
        except Exception as e:
            print(f"An error occurred: {e}")
            df_filled = pd.DataFrame()  # Return an empty DataFrame in case of error

        return df_filled  

    @staticmethod
    def fill_missing_values_linear(df):
        df_filled = df.copy()
        initial_imputer = SimpleImputer(strategy='mean')
        df_initial_imputed = pd.DataFrame(initial_imputer.fit_transform(df_filled), columns=df.columns)
        for column in df.columns:
            missing_values_mask = df[column].isnull()
            if not missing_values_mask.any():
                continue
            X_train = df_initial_imputed.loc[~missing_values_mask].drop(columns=[column])
            y_train = df_initial_imputed.loc[~missing_values_mask, column]
            X_test = df_initial_imputed.loc[missing_values_mask].drop(columns=[column])
            model = LinearRegression()
            model.fit(X_train, y_train)
            predicted_values = model.predict(X_test)
            df_filled.loc[missing_values_mask, column] = predicted_values
        return df_filled

    @staticmethod
    def fill_missing_values_rf(df):
        df_filled = df.copy()
        initial_imputer = SimpleImputer(strategy='mean')
        df_initial_imputed = pd.DataFrame(initial_imputer.fit_transform(df_filled), columns=df.columns)
        for column in df.columns:
            missing_values_mask = df[column].isnull()
            if not missing_values_mask.any():
                continue
            X_train = df_initial_imputed.loc[~missing_values_mask].drop(columns=[column])
            y_train = df_initial_imputed.loc[~missing_values_mask, column]
            X_test = df_initial_imputed.loc[missing_values_mask].drop(columns=[column])
            model = RandomForestRegressor()
            model.fit(X_train, y_train)
            predicted_values = model.predict(X_test)
            df_filled.loc[missing_values_mask, column] = predicted_values
        return df_filled

    @staticmethod
    def poz_transform(df, method='standardize'):
        try:
            numerical_columns = df.select_dtypes(include='number')
            if method == 'standardize':
                transformed_df = (numerical_columns - numerical_columns.mean()) / numerical_columns.std()
            elif method == 'boxcox':
                transformed_df = numerical_columns.apply(lambda x: pd.Series(boxcox(x + 1)[0], index=x.index))
            elif method == 'yeojohnson':
                pt = PowerTransformer(method='yeo-johnson')
                transformed_data = pt.fit_transform(numerical_columns)
                transformed_df = pd.DataFrame(transformed_data, columns=numerical_columns.columns, index=numerical_columns.index)
            elif method == 'normalize':
                transformed_df = numerical_columns.apply(lambda x: (x - x.min()) / (x.max() - x.min()))
            elif method == 'scaler':
                scaler = StandardScaler()
                transformed_data = scaler.fit_transform(numerical_columns)
                transformed_df = pd.DataFrame(transformed_data, columns=numerical_columns.columns, index=numerical_columns.index)
            elif method == 'minmax':
                scaler = MinMaxScaler()
                transformed_data = scaler.fit_transform(numerical_columns)
                transformed_df = pd.DataFrame(transformed_data, columns=numerical_columns.columns, index=numerical_columns.index)
            elif method == 'log2':
                transformed_df = np.log2(numerical_columns)
            else:
                raise ValueError(
                    "Invalid method. Options: 'standardize', 'boxcox', 'yeojohnson', 'normalize', 'scaler', 'minmax', 'log2'.")

            # Combine transformed numerical columns with non-numeric columns
            for column in df.columns:
                if column in numerical_columns.columns:
                    df[column] = transformed_df[column]
        except Exception as e:
            print(f"An error occurred: {e}")
            df = pd.DataFrame()  # Return an empty DataFrame in case of error
        return df
    
    def poz_outliers(data, method='z-score', threshold=3, action='remove'):
        cleaned_data = data.copy()
        if method == 'z-score':
            # Detect outliers using z-score
            z_scores = (cleaned_data - cleaned_data.mean()) / cleaned_data.std()
            outliers = abs(z_scores) > threshold

        elif method == 'iqr':
            # Detect outliers using interquartile range (IQR)
            q1 = cleaned_data.quantile(0.25)
            q3 = cleaned_data.quantile(0.75)
            iqr = q3 - q1
            lower_bound = q1 - threshold * iqr
            upper_bound = q3 + threshold * iqr
            outliers = (cleaned_data < lower_bound) | (cleaned_data > upper_bound)

        else:
            raise ValueError("Invalid method. Options: 'z-score', 'iqr'")

        if action == 'remove':
            # Remove outliers
            cleaned_data = cleaned_data[~outliers.any(axis=1)]

        elif action == 'cap':
            # Cap outliers to a specified range
            cleaned_data[outliers] = cleaned_data.clip(lower=lower_bound, upper=upper_bound, axis=1)

        else:
            raise ValueError("Invalid action. Options: 'remove', 'transform', 'cap'")

        return cleaned_data
    
    @staticmethod
    def poz_data_aggregation(data, target_column):
        """
        Fill missing values in the target column with the mean of that column.

        Parameters:
        - data: DataFrame containing the dataset.
        - target_column: The column to fill missing values with its mean.

        Returns:
        - DataFrame with filled missing values in the target column.
        """
        try:
            # Fill missing values in the target column with the mean of that column
            if target_column in data.columns and pd.api.types.is_numeric_dtype(data[target_column]):
                mean_value = data[target_column].mean()
                data[target_column].fillna(mean_value, inplace=True)
            
            return data
        except Exception as e:
            print(f"An error occurred during data aggregation: {e}")
            return None
    
    @staticmethod
    def poz_split_data(data, target_column, test_size, random_state, axis):
            """
            Split data into train and test sets.
            Parameters:
            - data: DataFrame containing the dataset.
            - target_column: Name of the target column.
            - test_size: Size of the test set.
            - random_state: Random state for reproducibility.
            - axis: Axis along which to drop the target column (0 for rows, 1 for columns).
            Returns:
            - X_train, X_test, y_train, y_test: Train and test sets for features and target.
            """
            if axis not in [0, 1]:
                raise ValueError("Axis must be 0 or 1.")

            X = data.drop(target_column, axis=axis)
            y = data[target_column]
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=random_state)
            return X_train, X_test, y_train, y_test
   
    
    @staticmethod
    def poz_feature_scaling(df, scaling_method='normalization'):
     try:
        df_scaled = df.copy()
        
        if scaling_method == 'normalization':
            df_scaled = df_scaled.apply(lambda x: (x - x.min()) / (x.max() - x.min()), axis=0)
        elif scaling_method == 'standard':
            scaler = StandardScaler()
            df_scaled = pd.DataFrame(scaler.fit_transform(df_scaled), columns=df_scaled.columns)
        elif scaling_method == 'cliplog':
            df_scaled = df_scaled.apply(lambda x: np.log1p(np.clip(x, a_min=1e-9, a_max=None)), axis=0)
        elif scaling_method == 'z-score':
            df_scaled = (df_scaled - df_scaled.mean()) / df_scaled.std()
        else:
            raise ValueError("Invalid scaling_method. Options: 'normalization', 'standard', 'cliplog', 'z-score'.")
        
     except Exception as e:
        print(f"An error occurred during feature scaling: {e}")
        df_scaled = pd.DataFrame()  # Return an empty DataFrame in case of error
    
     return df_scaled
     
    @staticmethod
    def feature_extraction(X, feature_names, y=None, method='pCA', n_components=2, k=5):
        # Impute missing values
        imputer = SimpleImputer(strategy='mean')
        X = imputer.fit_transform(X)
        
        if method == 'pCA':
            pca = PCA(n_components=n_components)
            X_transformed = pca.fit_transform(X)
            columns = [f"PC{i + 1}" for i in range(n_components)]
        elif method == 'scaling':
            scaler = MinMaxScaler()
            X_transformed = scaler.fit_transform(X)
            columns = feature_names
        elif method == 'polynomial':
            poly = PolynomialFeatures(degree=2, include_bias=False)
            X_transformed = poly.fit_transform(X)
            poly_feature_names = poly.get_feature_names_out(input_features=feature_names)
            columns = [name.replace(" ", "*") for name in poly_feature_names]
        elif method == 'select k-best':
            selector = SelectKBest(score_func=f_classif, k=k)
            X_transformed = selector.fit_transform(X, y)
            selected_indices = selector.get_support(indices=True)
            columns = [feature_names[i] for i in selected_indices]
        elif method == 'expansion':
            poly = PolynomialFeatures(degree=3, include_bias=False)
            X_transformed = poly.fit_transform(X)
            poly_feature_names = poly.get_feature_names_out(input_features=feature_names)
            columns = [name.replace(" ", "*") for name in poly_feature_names]
        else:
            raise ValueError(
                "Invalid method. Choose one of: 'pCA', 'scaling', 'polynomial', 'select k-best', 'expansion'")
        return pd.DataFrame(X_transformed, columns=columns)
    
    @staticmethod
    def plot_feature_importances(X, y, feature_names, algorithm='Random Forest', n_estimators_rf=100, n_estimators_xgb=100, random_state=42):
        if algorithm == 'Random Forest':
            clf = RandomForestClassifier(n_estimators=n_estimators_rf, random_state=random_state)
        elif algorithm == 'XGBoost':
            clf = xgb.XGBClassifier(n_estimators=n_estimators_xgb, random_state=random_state, use_label_encoder=False)
        else:
            return "Invalid algorithm selected.", None

        clf.fit(X, y)
        importances = clf.feature_importances_
        indices = np.argsort(importances)[::-1]

        feature_ranking = [(f + 1, feature_names[indices[f]], float(importances[indices[f]])) for f in range(X.shape[1])]  # Convert float32 to float

        plt.figure(figsize=(10, 6))
        plt.title("Feature Importances")
        plt.bar(range(X.shape[1]), importances[indices], align="center")
        plt.xticks(range(X.shape[1]), [feature_names[i] for i in indices], rotation=90)
        plt.xlabel("Feature")
        plt.ylabel("Feature Importance")

        # Save plot to a BytesIO object and encode it to base64
        img = io.BytesIO()
        plt.savefig(img, format='png')
        img.seek(0)
        plot_url = base64.b64encode(img.getvalue()).decode('utf8')
        plt.close()

        return feature_ranking, plot_url

        
    
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return 'No file part'
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'
    df = pd.read_csv(file)
    feature_names = df.columns[:-1]
    target_column = df.columns[-1]

    X = df[feature_names].values
    y = df[target_column].values

    algorithm = request.form.get('algorithm', 'Random Forest')
    n_estimators_rf = int(request.form.get('n_estimators_rf', 100))
    n_estimators_xgb = int(request.form.get('n_estimators_xgb', 100))

    feature_ranking, plot_url = preprocess.plot_feature_importances(X, y, feature_names, algorithm, n_estimators_rf, n_estimators_xgb)

    return render_template('result.html', feature_ranking=feature_ranking, plot_url=plot_url)

@app.route('/feature-importance', methods=['POST'])
def feature_importance():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"})
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"})

    df = pd.read_csv(file)
    target_column = request.form.get('targetColumn')
    algorithm = request.form.get('algorithm', 'Random Forest')
    n_estimators_rf = int(request.form.get('n_estimators_rf', 100))
    n_estimators_xgb = int(request.form.get('n_estimators_xgb', 100))

    # Ensure there are no NaN values in the target column
    df = df.dropna(subset=[target_column])

    # Convert the target column to integers if necessary
    df[target_column] = df[target_column].astype(int)

    feature_names = df.columns[df.columns != target_column]
    X = df[feature_names].values
    y = df[target_column].values

    feature_ranking, plot_url = preprocess.plot_feature_importances(X, y, feature_names, algorithm, n_estimators_rf, n_estimators_xgb)

    # Prepare response
    feature_ranking_json = [{"rank": rank, "feature_name": feature_name, "importance": float(importance)} for rank, feature_name, importance in feature_ranking]

    return jsonify({"feature_ranking": feature_ranking_json, "plot_url": plot_url})


@app.route('/feature-extraction', methods=['POST'])
def extract_features():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    method = request.form.get('method', 'pCA')
    n_components = int(request.form.get('n_components', 2))
    k = int(request.form.get('k', 5))

    df = pd.read_csv(file)
    X = df.iloc[:, :-1].values  # Assuming the last column is the target
    feature_names = df.columns[:-1].tolist()
    y = df.iloc[:, -1].values if method == 'select k-best' else None

    try:
        result = preprocess.feature_extraction(X, feature_names, y, method, n_components, k)
        return result.to_json(orient='split')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400


 
# Endpoint to handle missing values
@app.route('/handle_missing_values', methods=['POST'])
def handle_missing_values():
    try:
        file = request.files['file']
        fill_method = request.form.get('method', 'mean')
        constant = request.form.get('constant', 0)
        df = preprocess.poz_read_file(file)

        if df is None:
            raise ValueError("Error reading the file. Unsupported file type or corrupt file.")

        df_filled = preprocess.poz_handle_missing_values(df, fill_method, constant)
        result = df_filled.to_json(orient="split")
        print(json.loads(result))
        return jsonify({'processed_data': json.loads(result)}), 200
        
    except Exception as e:
        print(f"Error processing the DataFrame: {e}")
        return jsonify({'error': f"Error processing the DataFrame: {e}"}), 500

# Endpoint to handle data transformation
@app.route('/transform_data', methods=['POST'])
def transform_data():
    try:
        file = request.files['file']
        method = request.form.get('method', 'standardize')
        df = preprocess.poz_read_file(file)

        if df is None:
            raise ValueError("Error reading the file. Unsupported file type or corrupt file.")

        df_transformed = preprocess.poz_transform(df, method)
        result = df_transformed.to_json(orient="split")
        print(json.loads(result))
        return jsonify({'transformed_data': json.loads(result)}), 200
        
    except Exception as e:
        print(f"Error transforming the DataFrame: {e}")
        return jsonify({'error': f"Error transforming the DataFrame: {e}"}), 500

# Endpoint to upload CSV file
@app.route('/upload-csv', methods=['POST'])
def upload_csv():
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        try:
            # Read CSV file using pandas
            df = pd.read_csv(file)

            # Convert DataFrame to JSON
            json_data = df.to_json(orient='split')

            # Return JSON response
            return jsonify({'data': json_data})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        
@app.route('/aggregate_data', methods=['POST'])
def handle_aggregate_data():
    try:
        # Check if the POST request has the file part
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file = request.files['file']
        target_column = request.form.get('targetColumn', None)

        # Load CSV data from file
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if file and target_column:
            # Read CSV data
            csv_string = file.stream.read().decode('utf-8')
            csv_data = StringIO(csv_string)

            # Load CSV into DataFrame
            df = pd.read_csv(csv_data)

            # Perform data aggregation using the predefined function
            aggregated_data = preprocess.poz_data_aggregation(df, target_column)

            if aggregated_data is not None:
                # Convert aggregated data to JSON format
                # aggregated_data_json = aggregated_data.to_json(orient='records')
                # return jsonify({'aggregated_data': aggregated_data_json}), 200
                aggregated_data_json = aggregated_data.to_json(orient='split')
                return jsonify({'aggregated_data':json.loads(aggregated_data_json)}), 200
        #      result = df_scaled.to_json(orient="split")
        # return jsonify({'scaled_data': json.loads(result)}), 200
            else:
                return jsonify({'error': 'Data aggregation failed'}), 500

        return jsonify({'error': 'Missing target column'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

        
        # Endpoint to handle outliers
@app.route('/handle_outliers', methods=['POST'])
def handle_outliers():
    try:
        file = request.files['file']
        method = request.form.get('method', 'z-score')
        threshold = float(request.form.get('threshold', 3))
        action = request.form.get('action', 'remove')
        df = preprocess.poz_read_file(file)

        if df is None:
            raise ValueError("Error reading the file. Unsupported file type or corrupt file.")

        df_outliers_handled = preprocess.poz_outliers(df, method, threshold, action)
        result = df_outliers_handled.to_json(orient="split")
        print(json.loads(result))
        return jsonify({'processed_data': json.loads(result)}), 200
        
    except Exception as e:
        print(f"Error processing outliers: {e}")
        return jsonify({'error': f"Error processing outliers: {e}"}), 500
    
        # Endpoint to handle data segmentation
@app.route('/segment_data', methods=['POST'])
def segment_data():
    try:
        # Check if file is present in the request
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']

        # Read the CSV file
        segmented_data = preprocess.poz_read_file(file)

        # Get the last column name as the target column
        target_column = segmented_data.columns[-1]

        # Split the data
        X_train, X_test, y_train, y_test = preprocess.poz_split_data(segmented_data, target_column=target_column, test_size=0.2, random_state=42, axis=1)

        # Convert split data to JSON
        split_data_json = {
            'X_train': X_train.to_json(orient='split'),
            'X_test': X_test.to_json(orient='split'),
            'y_train': y_train.to_frame(name=target_column).to_json(orient='split'),
            'y_test': y_test.to_frame(name=target_column).to_json(orient='split'),
            'X_train_count':len(X_train),
            'X_test_count':len(X_test),
            'y_train_count': len(y_train),
            'y_test_count': len(y_test)
        }

        return jsonify({'split_data': split_data_json}), 200
    except Exception as e:
        print(f"Error segmenting the data: {e}")
        return jsonify({'error': f"Error segmenting the data: {e}"}), 500
    
    # Scale features in the dataset

# @app.route('/scale', methods=['POST'])
# def scale():
#     try:
#         data = request.get_json()
#         df = pd.DataFrame(data['data'])
#         scaling_method = data.get('scaling_method', 'normalization')
#         df_scaled = preprocess.poz_feature_scaling(df, scaling_method)
#         return jsonify(df_scaled.to_dict(orient='records'))
#     except Exception as e:
#         return jsonify({'error': str(e)}), 400
@app.route('/scale', methods=['POST'])
def scale():
    try:
        # Retrieve CSV file from FormData
        file = request.files['file']
        
        # Read CSV file into pandas DataFrame
        df = pd.read_csv(file)
        
        # Retrieve scaling method from FormData
        scaling_method = request.form.get('scaling_method', 'normalization')
        
        # Perform feature scaling based on the chosen method
        df_scaled = preprocess.poz_feature_scaling(df, scaling_method)
        
        # Convert scaled DataFrame to JSON and return
        # return jsonify(df_scaled.to_json(orient='records'))
        result = df_scaled.to_json(orient="split")
        return jsonify({'scaled_data': json.loads(result)}), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 400

if __name__ == '__main__':
    app.run(debug=True)














