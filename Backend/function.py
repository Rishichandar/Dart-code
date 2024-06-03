# import os
# import pandas as pd
# import docx
# import io
# import json
# class Function:
#     def poz_read_file(file_path_or_uploaded_file):
#         # Check if the input is a file path or an uploaded file object
#         if isinstance(file_path_or_uploaded_file, str):
#             # It's a file path
#             file_extension = os.path.splitext(file_path_or_uploaded_file)[-1].lower()
#             with open(file_path_or_uploaded_file, 'rb') as f:
#                 file_content = f.read()
#         else:
#             # It's an uploaded file object
#             file_extension = file_path_or_uploaded_file.name.split('.')[-1].lower()
#             file_content = file_path_or_uploaded_file.getvalue()

#         # Read the file content based on its extension
#         if file_extension == 'csv':
#             return pd.read_csv(io.BytesIO(file_content))
#         elif file_extension in ['xls', 'xlsx']:
#             return pd.read_excel(io.BytesIO(file_content))
#         elif file_extension == 'txt':
#             data = file_content.decode("utf-8")
#             return pd.DataFrame({'text_data': [data]})
#         elif file_extension in ['doc', 'docx']:
#             doc = docx.Document(io.BytesIO(file_content))
#             text = ' '.join([para.text for para in doc.paragraphs])
#             return pd.DataFrame({'text_data': [text]})
#         elif file_extension == 'json':
#             json_data = json.loads(file_content)
#             text = ' '.join(str(value) for value in json_data.values())
#             return pd.DataFrame({'text_data': [text]})
#         else:
#             print("Unsupported file type.")
#             return None
from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import pandas as pd
import docx
import io
import json
import warnings



app = Flask(__name__)
CORS(app)  # Enable CORS

class Function:
    @staticmethod
    def poz_read_file(file_path_or_uploaded_file):
        if isinstance(file_path_or_uploaded_file, str):
            file_extension = os.path.splitext(file_path_or_uploaded_file)[-1].lower()
            with open(file_path_or_uploaded_file, 'rb') as f:
                file_content = f.read()
        else:
            file_extension = file_path_or_uploaded_file.filename.split('.')[-1].lower()
            file_content = file_path_or_uploaded_file.read()

        if file_extension == 'csv':
            return pd.read_csv(io.BytesIO(file_content)).to_dict(orient='records')
        elif file_extension in ['xls', 'xlsx']:
            return pd.read_excel(io.BytesIO(file_content)).to_dict(orient='records')
        elif file_extension == 'txt':
            data = file_content.decode("utf-8")
            return {'text_data': data}
        elif file_extension in ['doc', 'docx']:
            doc = docx.Document(io.BytesIO(file_content))
            text = ' '.join([para.text for para in doc.paragraphs])
            return {'text_data': text}
        elif file_extension == 'json':
            json_data = json.loads(file_content)
            return json_data
        else:
            print("Unsupported file type.")
            return None

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    data = Function.poz_read_file(file)
    if data is None:
        return jsonify({'error': 'Unsupported file type'}), 400

    return jsonify(data)

if __name__ == '__main__':
    app.run(debug=True)
